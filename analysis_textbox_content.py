# -*- coding: UTF-8 -*-

import docx
import io
import json

word_path = "/Users/wuziyue/Downloads/数据中台_底层数据方案.docx"
word_path = "./docx_data/wuliangye.docx"


def test1():
    document = docx.Document(word_path)
    # part = document.part

    print(len(document.paragraphs))
    for para in document.paragraphs:
        print(para.style.name + ": " + para.text)

    print(len(document.sections))
    for section in document.sections:
        print(section.header)
        print(section.footer)

    print(len(document.tables))
    for table in document.tables:
        for i in range(1, len(table.rows)):
            row_text_list = []
            for j in range(len(table.columns)):
                row_text_list.append(table.cell(i, j).text)
            print(' '.join(row_text_list))

    print(len(document.inline_shapes))
    for inline_shape in document.inline_shapes:
        # type docx.enum.shape.WD_INLINE_SHAPE
        # CHART = 12
        #     LINKED_PICTURE = 4
        #     PICTURE = 3
        #     SMART_ART = 15
        #     NOT_IMPLEMENTED = -6
        print("type: %s, width: %f, height: %f" % (inline_shape.type, inline_shape.width.cm, inline_shape.height.cm))

    print(document.settings.element)

    print(document.styles.default)


def test2():
    # 提取word文本框中文本
    document = docx.Document(word_path)
    children = document.element.body.iter()

    textbox_list = []
    tags = []
    for child in children:
        # 通过类型判断目录
        child_iters = []
        if child.tag.endswith(('AlternateContent', 'textbox')):
            for ci in child.iter():
                tags.append(ci.tag)
                if ci.tag.endswith(('main}r', 'main}pPr')):
                    child_iters.append(ci)
            if len(child_iters) > 0:
                textbox_list.append(child_iters)

    textbox_text_list = []
    for textbox in textbox_list:
        text = []
        for ci in textbox:
            print(ci)
            if ci.tag.endswith('main}pPr'):
                text.append('')
            else:
                text[-1] += ci.text
            ci.text = ''
            print(ci.attrib)
        if len(text) > 0:
            textbox_text_list.append(text)

    for text_list in textbox_text_list:
        # for text in text_list:
        #     print(text)
        print("------------------")
        print(text_list)


def test3():
    # 提取word文本框中文本
    document = docx.Document(word_path)
    children = document.element.body.iter()
    for child in children:
        # sub_list = child.getchildren()
        # if len(sub_list) == 0:
        #     continue
        print("1----------------------")
        print("tag: " + child.tag)
        print("text: " + get_string(child.text))
        print("attrib: " + get_string(child.attrib))
        # for sub in sub_list:
        #     sub_list_2 = sub.getchildren()
        #     if len(sub_list_2) == 0:
        #         continue
        #     print("2--------")
        #     print("tag: " + sub.tag)
        #     print("text: " + get_string(sub.text))
        #     print("attrib: " + get_string(sub.attrib))
        #     for sub_2 in sub_list_2:
        #         sub_list_3 = sub.getchildren()
        #         if len(sub_list_3) == 0:
        #             continue
        #         print("3-")
        #         print("tag: " + sub_2.tag)
        #         print("text: " + get_string(sub_2.text))
        #         print("attrib: " + get_string(sub_2.attrib))


def get_string(text):
    if text is None:
        text = ""
    if type(text) is not str:
        text = str(text)
    return text


def test4():
    shape = ''
    shape_list = []
    document = docx.Document(word_path)
    children = document.element.body.iter()
    textbox_list = []

    for child in children:
        if child.tag.endswith('shape'):
            shape = child.attrib
        if child.tag.endswith('textbox'):
            sub_list = []
            for item in child.iter():
                sub_list.append(item)
            if len(sub_list) > 0:
                textbox_list.append(sub_list)
            shape_list.append(shape)

    print(len(textbox_list))
    print(len(shape_list))
    # print(textbox_list)

    total_dict_list = []
    for i in range(len(textbox_list)):
        textbox_dict_list = []
        temp_text = ""
        temp_font = ""
        temp_font_size = ""
        for child in textbox_list[i]:
            # print(child)
            tag = child.tag
            # print(tag)
            text = get_string(child.text)
            # print(text)
            attrib = get_string(child.attrib)

            if tag.endswith("main}r") and len(text) > 0:
                if len(temp_text) > 0:
                    print(temp_text)
                    textbox_dict_list.append({"text": temp_text,
                                              "font": temp_font,
                                              "font_size": temp_font_size})
                    temp_font = ""
                    temp_font_size = ""
                temp_text = text
                print(temp_text)

            if tag.endswith("main}rFonts") and len(attrib) > 0:
                temp_font = child.attrib.values()[0]
            if tag.endswith("main}sz") and len(attrib) > 0:
                temp_font_size = child.attrib.values()[0]
        total_dict_list.append(textbox_dict_list)
    save_json_('./docx_data/docx_analysis_result.txt', total_dict_list,shape_list)


def fix_coordinate(file_path):
    font_18_x = 10.5
    font_18_y = 10.5

    with io.open(file_path,"r",encoding="utf-8") as f:
        docx_results = f.readlines()
        list_of_textbox = []
        textbox_info = []
        for docx_result in docx_results:
            docx_dict = json.loads(docx_result)
            textbox_content = docx_dict['textbox_content']
            textbox_shape = docx_dict['textbox_shape']
            textbox_width = textbox_shape['width']
            textbox_height = textbox_shape['height']
            textbox_x = textbox_shape['margin-left']
            textbox_y = textbox_shape['margin-top']
            # textbox_page = docx_dict['page']
            textbox_page = 1
            textbox_text = ''

            textbox_info = {"position":[float(textbox_x.strip('pt')),float(textbox_y.strip('pt'))],
                            "page":textbox_page
                            }

            position_x = float(textbox_x.strip('pt'))
            position_y = float(textbox_y.strip('pt'))
            chars_info = []
            for run_item in textbox_content:
                font_size = run_item['font_size']
                text = run_item['text']
                font = run_item['font']
                textbox_text = textbox_text + text

                if text == '\n':
                    position_x = float(textbox_x.strip('pt'))
                    position_y = position_y + font_18_y
                    continue

                for w in text:
                    position_x = position_x + font_18_x
                    chars_dict = {'char':w,
                                  'position':[position_x,position_y],
                                  'font_size':font_size,
                                  'font':font
                                  }
                    chars_info.append(chars_dict)
                    if position_x+font_18_x >= float(textbox_x.strip('pt')) + float(textbox_width.strip('pt')):
                        position_x = float(textbox_x.strip('pt'))
                        position_y = position_y + font_18_y

            textbox_info['chars_info'] = chars_info
            textbox_info['text'] = textbox_text

        list_of_textbox.append(textbox_info)
    print(list_of_textbox)







def save_json(path, items):
    with io.open(path, 'w', encoding='utf-8') as f:
        for item in items:
            f.write(json.dumps(item, ensure_ascii=False) + '\n')


def save_json_(path, item_1,item_2):
    with io.open(path, 'w', encoding='utf-8') as f:
        for i in range(len(item_1)):
            f.write(json.dumps(item_1[i], ensure_ascii=False) + '\n')
            # shape_list_item = [item_2[i]]
            # f.write(json.dumps(shape_list_item, ensure_ascii=False) + '\n')
            f.write(str(item_2[i]) + '\n')
            f.write('-------------------------------------------------------'+'\n')


if __name__ == '__main__':
    # test1()
    # test2()
    # test3()
    # test4()
    fix_coordinate("./docx_data/000858_1.txt")