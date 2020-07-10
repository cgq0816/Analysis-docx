from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
"""
   Generate a reference to each paragraph and table child within *parent*,
   in document order. Each returned value is an instance of either Table or
   Paragraph. *parent* would most commonly be a reference to a main
   Document object, but also works for a _Cell object, which itself can
   contain paragraphs and tables.
"""
def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):         #判断两个参数是否为同一类型，返回布尔型
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

"""
main function to extract tables
"""
def extract_tables(document):
    count = 0
    current_context=''
    #iterator the blocks in doc
    for block in iter_block_items(document):
        # print(block.text if isinstance(block, Paragraph) else '<table>')
        if isinstance(block, Paragraph):
            # print("------------------text--------------------")
            print("text:  " + block.text)
            print("style:  " , block.style)
            for r in block.runs:
                print("no_1",r.text)
                print("no_2",r.element)
                print("no_3",r.font)
                print("no_4",r.part)
        elif isinstance(block, Table):
            current_context=''
            for row in block.rows:
                row_data = []
                for cell in row.cells:
                    text_cell=''
                    for paragraph in cell.paragraphs:
                        text_cell += paragraph.text.strip()
                    if text_cell is '':
                        text_cell="NULL"
                    row_data.append(text_cell)
                print("|".join(row_data))
if __name__ == '__main__':
    document = Document('./docx_data/zishiyingxuexi.docx')
    extract_tables(document)

# import docx
# from docx.document import Document as _Document
# from docx.oxml.text.paragraph import CT_P
# from docx.oxml.table import CT_Tbl
# from docx.table import _Cell, Table, _Row
# from docx.text.paragraph import Paragraph
#
# doc = docx.Document('./docx_data/zishiyingxuexi.docx')
#
#
# def table_nested_parsing(cell, current_row, current_col):
#     for block in cell._element:
#         if isinstance(block, CT_P):
#             print(Paragraph(block, cell).text)
#         if isinstance(block, CT_Tbl):
#             block = Table(block, cell)
#             for row in range(len(block.rows)):
#                 for col in range(len(block.columns)):
#                     cell_table = block.cell(row, col)
#                     table_nested_parsing(cell_table, row, col)
#
#
# def doc_parsing(doc):
#     doc_list = []
#     for doc_part in doc.element.body:
#         if isinstance(doc_part, CT_P):
#             print(Paragraph(doc_part, doc).text)
#         if isinstance(doc_part, CT_Tbl):
#             tb1 = Table(doc_part, doc)
#             for row in range(len(tb1.rows)):
#                 for col in range(len(tb1.columns)):
#                     cell_table = tb1.cell(row, col)
#                     table_nested_parsing(cell_table, row, col)
#
#
# if __name__ == "__main__":
#     doc_parsing(doc)